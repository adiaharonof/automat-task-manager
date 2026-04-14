from docx import Document
from TaskManager import *


def create_tasks_document(tasks: TaskManager):
    word = Document()
    word.add_heading("My Tasks", 0)
    word.add_heading("Completed:", 1)
    for completed_task in tasks.get_completed_tasks():
        word.add_paragraph(f"• {completed_task.title} - {completed_task.description}")

    word.add_heading("In Progress:", 1)
    for uncompleted_task in tasks.get_uncompleted_tasks():
        word.add_paragraph(f"• {uncompleted_task.title} - {uncompleted_task.description}")

    word.add_heading("Overdue:", 1)
    for overdue_task in tasks.get_overdue_tasks():
        word.add_paragraph(f"• {overdue_task.title} - {overdue_task.description}")
    word.save("tasks.docx")
